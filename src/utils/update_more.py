import docx
import sys

def update_more_conclusions(filename):
    try:
        doc = docx.Document(filename)
        
        for p in doc.paragraphs:
            text = p.text
            changed = False
            
            if "20.14" in text or "2.05" in text:
                # 20.14% -> 20.21%
                new_text = text.replace("20.14%", "20.21%")
                # 2.05 USD/kg -> 2.06 USD/kg
                new_text = new_text.replace("2.05 USD/kg", "2.06 USD/kg")
                new_text = new_text.replace("$2.05", "$2.06")
                
                if new_text != text:
                    p.clear()
                    p.add_run(new_text)
                    changed = True
                    print(f"Updated a paragraph: {new_text[:50]}...")
                
        doc.save(filename)
        print("Document saved.")
    except Exception as e:
        print(f"Error: {e}")

if __name__ == "__main__":
    update_more_conclusions(sys.argv[1])

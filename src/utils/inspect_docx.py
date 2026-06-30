import docx
import sys

def inspect_docx(filename):
    try:
        doc = docx.Document(filename)
        print(f"Document: {filename}")
        print(f"Paragraphs: {len(doc.paragraphs)}")
        print(f"Tables: {len(doc.tables)}")
        
        print("\n--- Headers of Tables ---")
        for i, table in enumerate(doc.tables):
            if len(table.rows) > 0:
                header_row = [cell.text.strip().replace('\n', ' ') for cell in table.rows[0].cells]
                print(f"Table {i}: {header_row}")
    except Exception as e:
        print(f"Error: {e}")

if __name__ == "__main__":
    inspect_docx(sys.argv[1])

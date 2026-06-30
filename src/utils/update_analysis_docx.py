import docx
from docx.shared import Inches
import sys

def update_docx_with_results(filename):
    try:
        doc = docx.Document(filename)
        
        for p in doc.paragraphs:
            text = p.text
            
            # Update T-Test records
            if "Se ejecut un script Python sobre 10,063 registros limpios" in text or "Se ejecutó un script Python sobre 10,063 registros limpios" in text:
                new_text = text.replace("10,063", "40,588")
                p.clear()
                p.add_run(new_text)
            
            # Update T-Test Results
            if "Resultados: P-Value =" in text and "Media EE.UU:" in text:
                new_text = "Resultados: P-Value = 2.499e-239. Media EE.UU: 2.37 USD vs Media Europa: 2.03 USD."
                p.clear()
                p.add_run(new_text)
                
            # Update T-Test note
            if "Nota metodolgica: Los estadsticos de esta subseccin" in text or "Nota metodológica" in text:
                # We can update or remove this note since we ran on all 40,588 now.
                new_text = "Nota metodológica: Los estadísticos de esta subsección fueron calculados sobre el total de 40,588 registros limpios correspondientes al período 2016-2024, garantizando máxima representatividad poblacional."
                p.clear()
                p.add_run(new_text)
                
            # Update Regression Results
            if "Resultados: R-cuadrado = 0.0065" in text:
                new_text = "Resultados: R-cuadrado = 0.000067"
                p.clear()
                p.add_run(new_text)
                
            # Insert T-Test Image
            if "[CAPTURA REQUERIDA: Script EDA y T-Test]" in text:
                p.clear()
                run = p.add_run()
                run.add_picture('ttest_plot.png', width=Inches(5.0))
                
            # Remove Note T-test
            if "[Nota: Captura de pantalla de la terminal mostrando los clculos de varianza y el resultado del P-Value.]" in text or "[Nota: Captura de pantalla de la terminal" in text:
                p.clear()
                
            # Insert Regression Image
            if "[CAPTURA REQUERIDA: Resultados Regresin Lineal]" in text or "[CAPTURA REQUERIDA: Resultados Regresión Lineal]" in text:
                p.clear()
                run = p.add_run()
                run.add_picture('regression_plot.png', width=Inches(5.0))
                
            # Remove Note Regression
            if "[Nota: Captura de pantalla del resultado del R-cuadrado en Python.]" in text:
                p.clear()
                
        doc.save(filename)
        print("Document updated with real results and images!")
    except Exception as e:
        print(f"Error: {e}")

if __name__ == "__main__":
    update_docx_with_results(sys.argv[1])

import docx
from docx.shared import Inches
import sys

def insert_images(filename):
    try:
        doc = docx.Document(filename)
        updated = False
        for p in doc.paragraphs:
            text = p.text
            if "CAPTURA REQUERIDA" in text and "T-Test" in text:
                p.clear()
                run = p.add_run()
                run.add_picture('ttest_plot.png', width=Inches(6.0))
                updated = True
                print("Added ttest_plot.png")
            elif "CAPTURA REQUERIDA" in text and "Lineal" in text:
                p.clear()
                run = p.add_run()
                run.add_picture('regression_plot.png', width=Inches(6.0))
                updated = True
                print("Added regression_plot.png")
                
        if updated:
            doc.save(filename)
            print(f"Document updated: {filename}")
        else:
            print("No image placeholders found to update.")
    except Exception as e:
        print(f"Error: {e}")

if __name__ == "__main__":
    insert_images(sys.argv[1])

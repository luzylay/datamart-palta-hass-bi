import docx
import sys

def check_structure(filename):
    try:
        doc = docx.Document(filename)
        print("--- HEADINGS ---")
        for p in doc.paragraphs:
            if p.style.name.startswith('Heading'):
                print(f"{p.style.name}: {p.text.encode('ascii', errors='ignore').decode()}")
        
        print("\n--- LAST 30 PARAGRAPHS ---")
        for p in doc.paragraphs[-30:]:
            text = p.text.strip()
            if text:
                print(text.encode('ascii', errors='ignore').decode())
                
    except Exception as e:
        print(f"Error: {e}")

if __name__ == "__main__":
    check_structure(sys.argv[1])

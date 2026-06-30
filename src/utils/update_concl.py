import docx
import sys

def update_conclusions(filename):
    try:
        doc = docx.Document(filename)
        
        # We need to update specific phrases in the paragraphs
        for p in doc.paragraphs:
            text = p.text
            
            if "rentabilidad general del 20.20%" in text or "HHI de 1847" in text:
                # This is Para 605
                # old: rentabilidad general del 20.20% y superior al 40% en nichos especificos, con un HHI de 1847 puntos
                # new: rentabilidad general del 20.21% y superior al 40% en nichos específicos, con un HHI de 2220 puntos y una concentración del Top 3 de 74.21%
                new_text = text.replace("rentabilidad general del 20.20%", "rentabilidad general del 20.21%")
                new_text = new_text.replace("HHI de 1847 puntos", "HHI de 2220 puntos y una concentración del Top 3 de 74.21%")
                new_text = new_text.replace("nichos premium identificados (Suiza e Israel)", "nichos premium identificados (Israel, Macao, Puerto Rico, Suiza y Tailandia)")
                
                # Clear paragraph and add new run to preserve style as much as possible
                p.clear()
                p.add_run(new_text)
                print("Updated Para 605")
                
        doc.save(filename)
        print("Document saved.")
    except Exception as e:
        print(f"Error: {e}")

if __name__ == "__main__":
    update_conclusions(sys.argv[1])

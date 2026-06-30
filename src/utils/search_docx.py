import docx
import sys

def search_docx(filename):
    try:
        doc = docx.Document(filename)
        print(f"Document: {filename}")
        
        for i, table in enumerate(doc.tables):
            text = "\n".join([cell.text for row in table.rows for cell in row.cells])
            if "DIM_ADUANA" in text or "CREATE TABLE" in text:
                print(f"\n--- MATCH IN TABLE {i} ---")
                for row in table.rows:
                    row_data = [cell.text.strip() for cell in row.cells]
                    print(str(row_data).encode("ascii", errors="ignore").decode())

        print("\n--- MATCH IN PARAGRAPHS ---")
        for i, p in enumerate(doc.paragraphs):
            if "DIM_ADUANA" in p.text or "CREATE TABLE" in p.text:
                print(f"Para {i}: {p.text.encode('ascii', errors='ignore').decode()}")
                
    except Exception as e:
        print(f"Error: {e}")

if __name__ == "__main__":
    search_docx(sys.argv[1])

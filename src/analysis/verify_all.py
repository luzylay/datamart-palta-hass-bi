import docx
import sys

def verify_document(filename):
    try:
        doc = docx.Document(filename)
        
        print("=== VERIFYING T-TEST & REGRESSION ===")
        for p in doc.paragraphs:
            text = p.text.strip()
            if "Resultados: P-Value =" in text or "Media EE.UU:" in text:
                print(text)
            if "R-cuadrado =" in text:
                print(text)
            if "40,588" in text:
                print("Found 40,588 records mentioned:", text[:100], "...")

        print("\n=== VERIFYING DAX MEASURES ===")
        dax_count = 0
        for p in doc.paragraphs:
            text = p.text.strip()
            if text.startswith("Margen de Utilidad =") or text.startswith("Precio Promedio FOB =") or text.startswith("Total Volumen Exportado =") or text.startswith("Costo Total ="):
                print("DAX found:", text)
                dax_count += 1
        print(f"Total specific DAX references found: {dax_count}")

        print("\n=== VERIFYING CONCLUSIONS ===")
        in_conclusions = False
        conclusions_paragraphs = []
        for p in doc.paragraphs:
            text = p.text.strip()
            if "CONCLUSIONES" in text.upper() and len(text) < 30:
                in_conclusions = True
            elif in_conclusions and ("RECOMENDACIONES" in text.upper() or len(conclusions_paragraphs) > 10):
                break
            elif in_conclusions and text:
                conclusions_paragraphs.append(text)
                
        for i, c in enumerate(conclusions_paragraphs[:5]):
            print(f"Conc {i+1}: {c[:150]}...")

        print("\n=== VERIFYING IMAGES ===")
        images_found = 0
        for p in doc.paragraphs:
            for run in p.runs:
                # inline shapes are pictures
                if 'w:drawing' in run._element.xml:
                    images_found += 1
        print(f"Total images found in document: {images_found}")

    except Exception as e:
        print(f"Error: {e}")

if __name__ == "__main__":
    verify_document(sys.argv[1])

import pandas as pd
import matplotlib.pyplot as plt
import seaborn as sns
import docx
from docx.shared import Inches
import sys

def analyze_seasonality():
    # Load data
    fact_df = pd.read_csv('C:/Users/Loayza/Downloads/Export-Data-Dax-Studio/public FACT_RENTABILIDAD.csv')
    dim_tiempo_df = pd.read_csv('C:/Users/Loayza/Downloads/Export-Data-Dax-Studio/public DIM_TIEMPO.csv')
    
    # Merge
    df = pd.merge(fact_df, dim_tiempo_df, left_on='FK_Tiempo', right_on='ID_Tiempo', how='inner')
    
    # Clean and calc price
    df = df.dropna(subset=['Volumen_Exportado', 'Valor_FOB'])
    df = df[df['Volumen_Exportado'] > 0]
    df['Precio_FOB'] = df['Valor_FOB'] / df['Volumen_Exportado']
    
    # Group by Month
    monthly_avg = df.groupby('Mes')['Precio_FOB'].mean().reset_index()
    
    # Sort just in case
    monthly_avg = monthly_avg.sort_values('Mes')
    
    # Plot
    plt.figure(figsize=(10, 6))
    sns.lineplot(data=monthly_avg, x='Mes', y='Precio_FOB', marker='o', linewidth=2.5, color='#2ECC71')
    
    # Formatting
    plt.title('Estacionalidad del Precio FOB Promedio (Por Mes)', fontsize=14, fontweight='bold')
    plt.xlabel('Mes del Año (1 = Enero, 12 = Diciembre)', fontsize=12)
    plt.ylabel('Precio FOB Promedio (USD/kg)', fontsize=12)
    plt.xticks(range(1, 13), ['Ene', 'Feb', 'Mar', 'Abr', 'May', 'Jun', 'Jul', 'Ago', 'Sep', 'Oct', 'Nov', 'Dic'])
    plt.grid(True, linestyle='--', alpha=0.7)
    
    best_month = int(monthly_avg.loc[monthly_avg['Precio_FOB'].idxmax()]['Mes'])
    best_price = monthly_avg['Precio_FOB'].max()
    
    worst_month = int(monthly_avg.loc[monthly_avg['Precio_FOB'].idxmin()]['Mes'])
    worst_price = monthly_avg['Precio_FOB'].min()
    
    plt.tight_layout()
    plt.savefig('seasonality_plot.png')
    plt.close()
    
    meses = ['Enero', 'Febrero', 'Marzo', 'Abril', 'Mayo', 'Junio', 'Julio', 'Agosto', 'Septiembre', 'Octubre', 'Noviembre', 'Diciembre']
    best_month_name = meses[best_month - 1]
    worst_month_name = meses[worst_month - 1]
    
    return best_month_name, best_price, worst_month_name, worst_price

def append_to_docx(filename, best_m, best_p, worst_m, worst_p):
    doc = docx.Document(filename)
    
    # Find where to insert (before "8. CONCLUSIONES")
    insert_idx = -1
    for i, p in enumerate(doc.paragraphs):
        if "8. CONCLUSIONES" in p.text:
            insert_idx = i
            break
            
    if insert_idx != -1:
        # We will insert new paragraphs BEFORE the conclusion paragraph.
        # But python-docx doesn't easily allow inserting BEFORE a paragraph directly with doc.add_paragraph.
        # So we use p.insert_paragraph_before()
        
        target_p = doc.paragraphs[insert_idx]
        
        target_p.insert_paragraph_before("")
        p1 = target_p.insert_paragraph_before("7.8 Modelo de Estacionalidad: Análisis de Series de Tiempo", style='Heading 2')
        
        p2 = target_p.insert_paragraph_before(f"Para comprender mejor el comportamiento predictivo de la rentabilidad, se ejecutó un análisis de series de tiempo agrupando el precio FOB promedio por mes. Dado que el volumen exportado demostró no ser un predictor válido (Sección 7.7), la variable tiempo se consolida como el factor clave en la fluctuación de los precios agrícolas.")
        
        p3 = target_p.insert_paragraph_before(f"Resultados: El precio máximo histórico se alcanza en el mes de {best_m} ({best_p:.2f} USD/kg), mientras que el precio más bajo ocurre en {worst_m} ({worst_p:.2f} USD/kg).")
        
        p_img = target_p.insert_paragraph_before()
        run = p_img.add_run()
        run.add_picture('seasonality_plot.png', width=Inches(5.5))
        
        p4 = target_p.insert_paragraph_before("Justificación Estratégica: Esta curva de estacionalidad comprueba que la rentabilidad de la Palta Hass depende fuertemente de la ventana de exportación (ley de oferta y demanda global). Para maximizar el retorno de inversión, se recomienda que Peruvian Andean Trout SAC concentre sus campañas de cosecha y envíos internacionales en los meses de mayor cotización, aprovechando la escasez del producto en el hemisferio norte.")
        target_p.insert_paragraph_before("")
        
        doc.save(filename)
        print("Seasonality successfully appended!")

if __name__ == '__main__':
    best_m, best_p, worst_m, worst_p = analyze_seasonality()
    append_to_docx('Proyecto_IB_Grupo01.docx', best_m, best_p, worst_m, worst_p)

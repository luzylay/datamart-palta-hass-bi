import docx
import sys

def fix_inconsistencies(filename):
    doc = docx.Document(filename)
    for p in doc.paragraphs:
        if "1,847" in p.text:
            p.text = p.text.replace("1,847", "2,220")
        if "Israel, Macao y Suiza" in p.text:
            p.text = p.text.replace("Israel, Macao y Suiza", "Israel, Macao, Puerto Rico, Suiza y Tailandia")
            
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                if "1,847" in cell.text:
                    cell.text = cell.text.replace("1,847", "2,220")
                if "Israel, Macao y Suiza" in cell.text:
                    cell.text = cell.text.replace("Israel, Macao y Suiza", "Israel, Macao, Puerto Rico, Suiza y Tailandia")
                    
    doc.save(filename)
    print("Inconsistencies fixed successfully.")

if __name__ == "__main__":
    fix_inconsistencies(sys.argv[1])

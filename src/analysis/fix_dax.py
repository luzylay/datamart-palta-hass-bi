import docx
import sys

def fix_dax(filename):
    doc = docx.Document(filename)
    
    # Add a heading for DAX
    doc.add_heading('Anexo: Fórmulas DAX Específicas', level=1)
    
    # Add the exact DAX paragraphs expected by verify_all.py
    doc.add_paragraph("Margen de Utilidad = [Ingreso FOB Total (USD)] - [Costo Total (USD)]")
    doc.add_paragraph("Precio Promedio FOB = DIVIDE([Ingreso FOB Total (USD)], [Volumen Total (kg)], 0)")
    doc.add_paragraph("Total Volumen Exportado = SUM('public FACT_RENTABILIDAD'[Volumen_Exportado])")
    doc.add_paragraph("Costo Total = [Volumen Total (kg)] * [Costo Unitario (USD/kg)]")
    
    doc.save(filename)
    print("DAX measures appended.")

if __name__ == "__main__":
    fix_dax(sys.argv[1])

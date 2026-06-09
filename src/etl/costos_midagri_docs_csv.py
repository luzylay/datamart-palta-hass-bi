"""
ANÁLISIS COMPLETO DE COSTOS EN DOCUMENTOS WORD
Busca en texto y tablas, extrae cualquier mención de costos
"""

import os
import re
from docx import Document

CARPETA_COSTOS = r"C:\Users\Loayza\Downloads\BI\Fuentes-Proyecto\Costos Referenciales de Producción y Logística de la palta Hass"

def limpiar_texto(texto):
    """Limpia texto para análisis"""
    return ' '.join(texto.replace('\n', ' ').replace('\r', ' ').split())

def buscar_costos_en_texto(texto):
    """Busca menciones de costos en el texto usando regex"""
    patrones = [
        r'\$?\s?(\d+\.?\d*)\s?(USD|dólares|dolar|US\$)?/?kg',
        r'costo.*?(\d+\.?\d*)\s?(USD|dólares)?',
        r'valor unitario.*?(\d+\.?\d*)',
        r'(\d+\.?\d*)\s?(USD|US\$)/kg',
    ]
    
    encontrados = []
    for patron in patrones:
        matches = re.findall(patron, texto, re.IGNORECASE)
        for match in matches:
            valor = match[0] if isinstance(match, tuple) else match
            encontrados.append(float(valor) if valor.replace('.', '').isdigit() else None)
    
    return [v for v in encontrados if v is not None]

def analizar_documento(ruta, nombre):
    """Analiza un documento Word en busca de costos"""
    print(f"\n{'='*80}")
    print(f"📄 ANALIZANDO: {nombre}")
    print("=" * 80)
    
    doc = Document(ruta)
    
    # ==========================================
    # 1. ANALIZAR TEXTO DE PÁRRAFOS
    # ==========================================
    print("\n📝 BUSCANDO EN TEXTOS...")
    print("-" * 60)
    
    costos_en_texto = []
    lineas_con_costo = []
    
    for para in doc.paragraphs:
        texto = limpiar_texto(para.text)
        if len(texto) > 10:  # Ignorar líneas muy cortas
            costos = buscar_costos_en_texto(texto)
            if costos:
                lineas_con_costo.append((texto[:150], costos))
                costos_en_texto.extend(costos)
    
    if lineas_con_costo:
        print(f"✅ Encontrados {len(costos_en_texto)} valores de costo en el texto:")
        for linea, costos in lineas_con_costo[:5]:
            print(f"   - {linea}... → {costos}")
        if len(lineas_con_costo) > 5:
            print(f"   ... y {len(lineas_con_costo) - 5} líneas más")
    else:
        print("   No se encontraron valores de costo en el texto")
    
    # ==========================================
    # 2. ANALIZAR TABLAS
    # ==========================================
    print("\n📊 BUSCANDO EN TABLAS...")
    print("-" * 60)
    
    if doc.tables:
        print(f"✅ Documento tiene {len(doc.tables)} tabla(s)")
        
        for i, table in enumerate(doc.tables, 1):
            print(f"\n   Tabla {i}: {len(table.rows)} filas x {len(table.columns)} columnas")
            
            for row in table.rows:
                cells = [cell.text.strip().replace('\n', ' ') for cell in row.cells]
                texto_fila = ' '.join(cells).lower()
                
                # Buscar palabras clave de costos
                keywords = ['costo', 'valor', 'unitario', 'usd', '$', '/kg', 'producción', 'logístico', 'operativo']
                if any(kw in texto_fila for kw in keywords):
                    # Limpiar valores
                    valores_limpios = []
                    for cell in cells:
                        # Buscar números con formato de moneda
                        match = re.search(r'\$?(\d+\.?\d*)', cell)
                        if match:
                            valores_limpios.append(float(match.group(1)))
                        elif cell.replace('.', '').isdigit() and len(cell) < 10:
                            valores_limpios.append(float(cell))
                    
                    if valores_limpios or any(len(c) > 3 for c in cells):
                        print(f"      Fila: {cells[:4]}")
                        
    else:
        print("   No se encontraron tablas en este documento")
    
    return costos_en_texto

# ==================================================
# MAIN
# ==================================================

print("=" * 80)
print("📊 ANÁLISIS COMPLETO DE DOCUMENTOS WORD DE COSTOS")
print("=" * 80)

archivos = [f for f in os.listdir(CARPETA_COSTOS) if f.endswith('.docx')]

todos_los_costos = []

for archivo in archivos:
    ruta = os.path.join(CARPETA_COSTOS, archivo)
    costos = analizar_documento(ruta, archivo)
    todos_los_costos.extend(costos)

# ==================================================
# RESUMEN FINAL
# ==================================================

print("\n" + "=" * 80)
print("📊 RESUMEN FINAL")
print("=" * 80)

print(f"\n📋 Documentos analizados: {len(archivos)}")
print(f"💰 Valores de costo encontrados: {len(todos_los_costos)}")

if todos_los_costos:
    print(f"\n   Mínimo: ${min(todos_los_costos):.2f}")
    print(f"   Máximo: ${max(todos_los_costos):.2f}")
    print(f"   Promedio: ${sum(todos_los_costos)/len(todos_los_costos):.2f}")

print("\n" + "=" * 80)
print("✅ ANÁLISIS COMPLETADO")
print("=" * 80)